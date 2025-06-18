import streamlit as st
import requests
from bs4 import BeautifulSoup
from langchain_openai import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.chains import RetrievalQA
import os
from docx import Document
from datetime import datetime
from dotenv import load_dotenv
from io import BytesIO
import time
import json
import re
from urllib.parse import quote_plus, urljoin
import logging
from typing import List, Dict, Tuple

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

AVAILABLE_MODELS = {
    "GPT-4o": "gpt-4o",
    "GPT-4o Mini": "gpt-4o-mini", 
    "GPT-4 Turbo": "gpt-4-turbo-preview",
    "GPT-3.5 Turbo": "gpt-3.5-turbo"
}

class WebScraper:
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

    def search_google(self, query: str, num_results = 10):
        """Search Google and return results"""
        try:
            search_url = f"https://www.google.com/search?q={quote_plus(query)}&num={num_results}"
            response = self.session.get(search_url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            results = []
            for result in soup.find_all('div', class_='g')[:num_results]:
                title_elem = result.find('h3')
                link_elem = result.find('a')
                snippet_elem = result.find('span', class_=['aCOpRe', 'st'])
                
                if title_elem and link_elem:
                    title = title_elem.get_text()
                    link = link_elem.get('href')
                    snippet = snippet_elem.get_text() if snippet_elem else ""
                    
                    if link.startswith('http'):
                        results.append({
                            'title': title,
                            'url': link,
                            'snippet': snippet
                        })
            
            return results
        except Exception as e:
            logger.error(f"Google search failed: {e}")
            return []
    
    def scrape_content(self, url: str) -> str:
        """Scrape content from a URL"""
        try:
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            for element in soup.find_all(['script', 'style', 'nav', 'footer', 'header', 'aside']):
                element.decompose()
            content_selectors = [
                'article', 'main', '.content', '.post-content', 
                '.entry-content', '.article-content', 'p'
            ]
            
            content = ""

            for selector in content_selectors:
                elements = soup.select(selector)
                if elements:
                    content = ' '.join([elem.get_text().strip() for elem in elements])
                    break
            
            if not content:
                content = soup.get_text()
        
            content = re.sub(r'\s+', ' ', content).strip()
            return content[:5000] 
            
        except Exception as e:
            logger.error(f"Failed to scrape {url}: {e}")
            return ""
    
    def comprehensive_search(self, company_name: str, industry: str) -> Dict[str, List[str]]:
        """Perform comprehensive search for company and industry data"""
        search_queries = {
            'company_overview': f"{company_name} company overview business model",
            'financial_info': f"{company_name} revenue financial performance annual report",
            'recent_news': f"{company_name} news 2024 2025 recent developments",
            'competitors': f"{company_name} competitors competitive analysis {industry}",
            'industry_trends': f"{industry} industry trends 2024 2025 market analysis",
            'challenges': f"{company_name} challenges problems issues {industry}",
            'technology': f"{company_name} technology digital transformation AI automation"
        }
        
        all_content = {}
        
        for category, query in search_queries.items():
            logger.info(f"Searching for: {category}")
            search_results = self.search_google(query, 5)
            
            category_content = []
            for result in search_results:
                content = self.scrape_content(result['url'])
                if content:
                    category_content.append({
                        'title': result['title'],
                        'url': result['url'],
                        'content': content
                    })
                time.sleep(1)  
            
            all_content[category] = category_content
        
        return all_content

class RAGSystem:
    
    def __init__(self, openai_api_key: str):
        self.embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        self.vectorstore = None
        self.qa_chain = None
    
    def build_vectorstore(self, scraped_data: Dict[str, List[str]]):
        """Build vector store from scraped data"""
        all_texts = []
        metadatas = []
         
        for category, content_list in scraped_data.items():
            for item in content_list:
                if item['content']:
                    chunks = self.text_splitter.split_text(item['content'])
                    for chunk in chunks:
                        all_texts.append(chunk)
                        metadatas.append({
                            'category': category,
                            'title': item['title'],
                            'url': item['url']
                        })
        
        if all_texts:
            self.vectorstore = FAISS.from_texts(
                all_texts, 
                self.embeddings, 
                metadatas=metadatas
            )
    
    def setup_qa_chain(self, llm):
        """Setup QA chain with the vector store"""
        if self.vectorstore:
            self.qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=self.vectorstore.as_retriever(search_kwargs={"k": 5}),
                return_source_documents=True
            )
    
    def query(self, question: str) -> str:
        """Query the RAG system"""
        if self.qa_chain:
            try:
                result = self.qa_chain({"query": question})
                return result['result']
            except Exception as e:
                logger.error(f"RAG query failed: {e}")
                return "Unable to retrieve information from the knowledge base"
        return "Knowledge base not formed"

class OptimizedMarketResearch:
    
    def __init__(self, openai_api_key: str, model_name: str = "gpt-4o"):
        self.scraper = WebScraper()
        self.rag = RAGSystem(openai_api_key)
        self.llm = ChatOpenAI(
            model=model_name,
            temperature=0.1,
            api_key=openai_api_key
        )
        self.scraped_data = None
    
    def initialize_knowledge_base(self, company_name: str, industry: str):
        if not self.scraped_data:
            self.scraped_data = self.scraper.comprehensive_search(company_name, industry)
            self.rag.build_vectorstore(self.scraped_data)
            self.rag.setup_qa_chain(self.llm)
    
    def chat_query(self, question: str) -> str:
        """Handle interactive chat queries"""
        if not self.rag.qa_chain:
            return "Knowledge base not initialized. Please generate a report first."
        
        return self.rag.query(question)
    
    def generate_research_report(self, company_name: str, industry: str) -> str:
        """Generate comprehensive research report using RAG + powerful LLM call"""
        
        # Step 1: Initialize knowledge base
        self.initialize_knowledge_base(company_name, industry)
        
        # Step 2: Extract structured information using RAG (MINIMAL LLM CALLS)
        research_questions = {
            'company_overview': f"What is {company_name}? Provide company background, history, and business model.",
            'financial_performance': f"What is {company_name}'s financial performance, revenue, and market position?",
            'products_services': f"What products and services does {company_name} offer?",
            'competitors': f"Who are {company_name}'s main competitors in the {industry} industry?",
            'industry_trends': f"What are the current trends and challenges in the {industry} industry?",
            'recent_developments': f"What are the recent news and developments about {company_name}?",
            'challenges': f"What are the main challenges and pain points facing {company_name}?",
            'technology_gaps': f"What technology gaps or digital transformation needs does {company_name} have?"
        }
        
        research_findings = {}
        for key, question in research_questions.items():
            research_findings[key] = self.rag.query(question)
        
        ai_use_cases = self._generate_ai_use_cases(company_name, industry, research_findings)
        final_report = self._generate_final_report(company_name, industry, research_findings, ai_use_cases)
        
        return final_report
    
    def _generate_ai_use_cases(self, company_name: str, industry: str, research_findings: Dict) -> str:
        
        prompt = f"""
        Based on the following research findings about {company_name} in the {industry} industry, generate 5 specific AI use cases that could address their challenges and opportunities.

        Research Findings:
        - Company Overview: {research_findings.get('company_overview', 'N/A')}
        - Challenges: {research_findings.get('challenges', 'N/A')}
        - Technology Gaps: {research_findings.get('technology_gaps', 'N/A')}
        - Industry Trends: {research_findings.get('industry_trends', 'N/A')}

        For each AI use case, provide:
        1. Problem Statement
        2. AI Solution Description
        3. Expected Benefits
        4. Implementation Complexity (Low/Medium/High)
        5. Estimated ROI Timeline
        6. Required Technologies

        Format as a structured list with clear headings.
        """
        
        try:
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            logger.error(f"AI use case generation failed: {e}")
            return "Unable to generate AI use cases due to processing error."
    
    def _generate_final_report(self, company_name: str, industry: str, research_findings: Dict, ai_use_cases: str) -> str:
        """Generate final comprehensive report (1 LLM CALL)"""
        
        prompt = f"""
        Create a comprehensive, executive-ready market research report for {company_name} in the {industry} industry.

        Structure the report with the following sections:
        Provide a concise overview of key findings and recommendations.

        # Company Overview
        {research_findings.get('company_overview', 'N/A')}

        # Financial Performance & Market Position
        {research_findings.get('financial_performance', 'N/A')}

        # Products & Services
        {research_findings.get('products_services', 'N/A')}

        # Competitive Landscape
        {research_findings.get('competitors', 'N/A')}

        # Industry Analysis
        {research_findings.get('industry_trends', 'N/A')}

        # Recent Developments
        {research_findings.get('recent_developments', 'N/A')}

        # Key Challenges Identified
        {research_findings.get('challenges', 'N/A')}

        # AI Use Cases & Recommendations
        {ai_use_cases}

        # Implementation Roadmap
        Provide a prioritized roadmap for implementing the AI solutions.

        # Expected ROI & Benefits
        Summarize the expected return on investment and key benefits.

        # Next Steps & Recommendations
        Provide specific, actionable next steps.

        Format this as a professional report suitable for C-level executives. Use clear headings and bullet points where appropriate.
        """
        
        try:
            response = self.llm.invoke(prompt)
            return response.content
        except Exception as e:
            logger.error(f"Final report generation failed: {e}")
            return f"Report generation encountered an error: {str(e)}"

def create_word_document(content: str, company_name: str) -> Document:
    doc = Document()
    
    title = doc.add_heading(f'AI Solutions Report for {company_name}', 0)
    title.alignment = 1  

    doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Company: {company_name}")
    doc.add_paragraph("")
      
    lines = content.split('\n')
    current_paragraph = ""
    
    for line in lines:
        line = line.strip()
        if not line:
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            continue
   
        if line.startswith('# '):
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            doc.add_paragraph(line[3:], style='List Number')
        else:
            if current_paragraph:
                current_paragraph += " " + line
            else:
                current_paragraph = line
    
    if current_paragraph:
        doc.add_paragraph(current_paragraph)
    
    return doc

def validate_inputs(company_name: str, industry: str) -> Tuple[bool, str]:
    """Validate user inputs"""
    if not company_name or not company_name.strip():
        return False, "Company name is required"
    
    if not industry or not industry.strip():
        return False, "Industry is required"
    
    if len(company_name.strip()) < 2:
        return False, "Company name must be at least 2 characters long"
    
    if len(industry.strip()) < 2:
        return False, "Industry must be at least 2 characters long"
    
    return True, ""

# Streamlit UI
def main():
    st.set_page_config(
        page_title="GENESIS",
        page_icon="🔍",
        layout="wide"
    )
    
    st.title("GENESIS - Generative AI for Strategy and Industry Solutions")
    st.markdown("""
    **Key Features:**
    - 🕷️ **Web Scraping**
    - 🧠 **RAG System**  
    - ⚡ **Multiple LLM Models**
    - 💬 **Interactive Chat**
    - 📄 **Comprehensive Reports**
    """)
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # Model selection
        selected_model_name = st.selectbox(
            "🤖 Select LLM Model",
            options=list(AVAILABLE_MODELS.keys()),
            index=0,
            help="Choose the AI model for analysis"
        )
        
        # Mode selection
        mode = st.radio(
            "📋 Analysis Mode",
            options=["Automatic Report Generation", "Interactive Chat"],
            help="Choose between generating a complete report or interactive Q&A"
        )
        
        st.markdown("---")
        st.markdown("""
        **Model Info:**
        - **GPT-4o**: Best performance, higher cost
        - **GPT-4o Mini**: Good balance of speed and quality
        - **GPT-4 Turbo**: Fast and capable
        - **GPT-3.5 Turbo**: Fastest, lowest cost
        """)
    
    # Input section
    col1, col2 = st.columns(2)
    
    with col1:
        company_name = st.text_input(
            "Company Name",
            placeholder="e.g., Tesla, Microsoft, Starbucks, Facebook",
            help="Enter the full company name"
        )
    
    with col2:
        industry = st.text_input(
            "Industry",
            placeholder="e.g., Electric Vehicles, Cloud Computing, Social Media",
            help="Enter the primary industry or sector"
        )
    
    # Initialize session state
    if 'research_system' not in st.session_state:
        st.session_state.research_system = None
    if 'knowledge_base_ready' not in st.session_state:
        st.session_state.knowledge_base_ready = False
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'current_company' not in st.session_state:
        st.session_state.current_company = ""
    if 'current_industry' not in st.session_state:
        st.session_state.current_industry = ""
    
    # Check if inputs have changed
    if (company_name != st.session_state.current_company or 
        industry != st.session_state.current_industry):
        st.session_state.knowledge_base_ready = False
        st.session_state.chat_history = []
        st.session_state.current_company = company_name
        st.session_state.current_industry = industry
    
    if mode == "Automatic Report Generation":
        # Automatic Report Generation Mode
        if st.button("🚀 Generate Report", type="primary"):
            # Validate inputs
            is_valid, error_message = validate_inputs(company_name, industry)
            
            if not is_valid:
                st.error(f"❌ {error_message}")
                return
            
            # Show progress and generate report
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                # Initialize the optimized research system with selected model
                selected_model = AVAILABLE_MODELS[selected_model_name]
                research_system = OptimizedMarketResearch(OPENAI_API_KEY, selected_model)
                
                status_text.text("🕷️ Scraping web data...")
                progress_bar.progress(20)
                
                status_text.text("🧠 Building RAG knowledge base...")
                progress_bar.progress(50)
                
                status_text.text("⚡ Generating analysis...")
                progress_bar.progress(80)
                
                # Generate the report
                with st.spinner("Processing..."):
                    report_content = research_system.generate_research_report(
                        company_name.strip(), 
                        industry.strip()
                    )
                    
                    # Store research system for potential chat use
                    st.session_state.research_system = research_system
                    st.session_state.knowledge_base_ready = True
                    
                    progress_bar.progress(90)
                    status_text.text("📝 Formatting report...")
                    
                    # Create Word document
                    doc = create_word_document(report_content, company_name.strip())
                    bio = BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    
                    progress_bar.progress(100)
                    status_text.text("✅ Report generation completed!")
                    
                    # Display results
                    st.success(f"Market research completed successfully using {selected_model_name}!")
                    
                    # Show report preview
                    st.subheader("📋 Report Preview")
                    
                    if not report_content:
                        st.warning("No content generated. Please check the inputs and try again.")
                        return
                    
                    # Truncate content for preview if it's too long
                    preview_content = report_content[:3000] + "..." if len(report_content) > 3000 else report_content
                    st.markdown(preview_content)
                    
                    if len(report_content) > 3000:
                        st.info("Preview truncated. Download the full report below or switch to Interactive Chat mode to ask specific questions.")
                    
                    # Download button
                    st.download_button(
                        label="📥 Download Full Report",
                        data=bio.getvalue(),
                        file_name=f"Detailed_Report_{company_name.replace(' ', '_').replace(',', '')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="Click to download the complete report"
                    )
                    
            except Exception as e:
                st.error(f"❌ An error occurred: {str(e)}")
                st.error("Please check your internet connection and API key, then try again.")
            
            finally:
                progress_bar.empty()
                status_text.empty()
    
    else:
        # Interactive Chat Mode
        st.subheader("💬 Interactive Chat")
        
        # Initialize knowledge base button
        if not st.session_state.knowledge_base_ready:
            if st.button("🧠 Initialize Knowledge Base", type="primary"):
                is_valid, error_message = validate_inputs(company_name, industry)
                
                if not is_valid:
                    st.error(f" {error_message}")
                    return
                
                with st.spinner("Initializing knowledge base..."):
                    try:
                        selected_model = AVAILABLE_MODELS[selected_model_name]
                        research_system = OptimizedMarketResearch(OPENAI_API_KEY, selected_model)
                        research_system.initialize_knowledge_base(company_name.strip(), industry.strip())
                        
                        st.session_state.research_system = research_system
                        st.session_state.knowledge_base_ready = True
                        
                        st.success(f"Knowledge base initialized for {company_name} using {selected_model_name}!")
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f" Failed to initialize knowledge base: {str(e)}")
        
        else:
            st.success(f"✅ Knowledge base ready for {company_name} in {industry}")
            
            # Display chat history
            for chat in st.session_state.chat_history:
                with st.chat_message("user"):
                    st.write(chat["question"])
                with st.chat_message("assistant"):
                    st.write(chat["answer"])
            
            # Chat input
            if question := st.chat_input("Ask a question about the company..."):
                with st.chat_message("user"):
                    st.write(question)
                
                with st.chat_message("assistant"):
                    with st.spinner("Thinking..."):
                        try:
                            answer = st.session_state.research_system.chat_query(question)
                            st.write(answer)
                            
                            st.session_state.chat_history.append({
                                "question": question,
                                "answer": answer
                            })
                            
                        except Exception as e:
                            error_msg = f"❌ Error processing question: {str(e)}"
                            st.write(error_msg)
                            st.session_state.chat_history.append({
                                "question": question,
                                "answer": error_msg
                            })
            
            # Suggested questions
            st.subheader("💡 Suggested Questions")
            suggested_questions = [
                f"What are the main challenges facing {company_name}?",
                f"Who are {company_name}'s biggest competitors?",
                f"What are the latest developments at {company_name}?",
                f"What AI solutions could benefit {company_name}?",
                f"What are the current trends in the {industry} industry?",
                f"What is {company_name}'s financial performance?"
            ]
            
            cols = st.columns(2)
            for i, q in enumerate(suggested_questions):
                with cols[i % 2]:
                    if st.button(q, key=f"suggested_{i}")
                        with st.chat_message("user"):
                            st.write(q)
                        
                        with st.chat_message("assistant"):
                            with st.spinner("Thinking..."):
                                try:
                                    answer = st.session_state.research_system.chat_query(q)
                                    st.write(answer)
                                    
                                    st.session_state.chat_history.append({
                                        "question": q,
                                        "answer": answer
                                    })
                                    st.rerun()
                                    
                                except Exception as e:
                                    error_msg = f"❌ Error: {str(e)}"
                                    st.write(error_msg)
            
            if st.button("🗑️ Clear Chat History"):
                st.session_state.chat_history = []
                st.rerun()

if __name__ == "__main__":
    main()
